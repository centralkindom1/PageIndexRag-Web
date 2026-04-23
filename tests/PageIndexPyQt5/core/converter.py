import os
import json
import csv
import re
from pathlib import Path
from typing import Optional

class Converter:
    @staticmethod
    def txt_to_markdown(file_path: str, output_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        filename = Path(file_path).stem
        lines = content.split("\n")
        md_lines = [f"# {filename}\n"]

        section_count = 0
        for line in lines:
            stripped = line.strip()
            is_header = False
            header_level = 2

            if stripped and len(stripped) > 2:
                if stripped.isupper() and len(stripped) > 3 and not stripped.startswith("---"):
                    is_header = True
                    header_level = 2
                elif re.match(r'^(\d+\.)+\s+\S', stripped):
                    dots = stripped.split()[0].count(".")
                    header_level = min(2 + dots - 1, 5)
                    is_header = True
                elif re.match(r'^第.+[章篇部]', stripped):
                    is_header = True
                    header_level = 2
                elif re.match(r'^第.+[节]', stripped):
                    is_header = True
                    header_level = 3

            if is_header:
                section_count += 1
                md_lines.append(f'\n{"#" * header_level} {stripped}\n')
            else:
                md_lines.append(line)

        if section_count == 0 and len(lines) > 100:
            md_lines = [f"# {filename}\n"]
            chunk_size = 50
            for i in range(0, len(lines), chunk_size):
                chunk = lines[i:i + chunk_size]
                section_num = i // chunk_size + 1
                md_lines.append(f"\n## Section {section_num}\n")
                md_lines.extend(chunk)

        result = "\n".join(md_lines)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(result)
        return output_path

    @staticmethod
    def json_to_markdown(file_path: str, output_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            data = json.load(f)

        filename = Path(file_path).stem
        md_lines = [f"# {filename}\n"]

        def _find_title_field(obj: dict) -> Optional[str]:
            for key in ["name", "title", "label", "id", "名称", "标题", "姓名"]:
                if key in obj and obj[key]:
                    return str(obj[key])[:80]
            return None

        def _format_json_value(v) -> str:
            if isinstance(v, (dict, list)):
                s = json.dumps(v, ensure_ascii=False)
                if len(s) > 200:
                    s = s[:197] + "..."
                return f"`{s}`"
            return str(v).replace("\n", " ")

        def _dict_to_md(data: dict, md_lines: list, level: int = 2):
            for key, value in data.items():
                if isinstance(value, dict):
                    md_lines.append(f'\n{"#" * min(level, 6)} {key}\n')
                    _dict_to_md(value, md_lines, level + 1)
                elif isinstance(value, list):
                    md_lines.append(f'\n{"#" * min(level, 6)} {key}\n')
                    if len(value) > 0 and isinstance(value[0], dict):
                        for i, item in enumerate(value[:50]):
                            title = _find_title_field(item) or f"Item {i + 1}"
                            md_lines.append(f'\n{"#" * min(level + 1, 6)} {title}\n')
                            for k, v in item.items():
                                md_lines.append(f"- **{k}**: {_format_json_value(v)}")
                    else:
                        for item in value[:200]:
                            md_lines.append(f"- {item}")
                else:
                    md_lines.append(f"- **{key}**: {_format_json_value(value)}")

        if isinstance(data, list) and len(data) > 0:
            if isinstance(data[0], dict):
                md_lines.append(f"\n## Data Overview\n")
                md_lines.append(f"Total records: {len(data)}\n")
                keys = list(data[0].keys())
                display_keys = keys[:8]
                md_lines.append("| " + " | ".join(str(k) for k in display_keys) + " |")
                md_lines.append("| " + " | ".join("---" for _ in display_keys) + " |")
                for item in data[:100]:
                    row = []
                    for k in display_keys:
                        val = str(item.get(k, "")).replace("|", "\\|").replace("\n", " ")
                        if len(val) > 60: val = val[:57] + "..."
                        row.append(val)
                    md_lines.append("| " + " | ".join(row) + " |")

                md_lines.append(f"\n## Record Details\n")
                for i, item in enumerate(data[:50]):
                    title = _find_title_field(item) or f"Record {i + 1}"
                    md_lines.append(f"\n### {title}\n")
                    for k, v in item.items():
                        md_lines.append(f"- **{k}**: {_format_json_value(v)}")
            else:
                md_lines.append(f"\n## Items ({len(data)} total)\n")
                for i, item in enumerate(data[:200]):
                    md_lines.append(f"{i + 1}. {item}")
        elif isinstance(data, dict):
            _dict_to_md(data, md_lines, level=2)
        else:
            md_lines.append(f"\n```json\n{json.dumps(data, ensure_ascii=False, indent=2)}\n```\n")

        result = "\n".join(md_lines)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(result)
        return output_path

    @staticmethod
    def csv_to_markdown(file_path: str, output_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            rows = list(reader)

        if not rows:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("# Empty CSV\n\nThis CSV file contains no data.\n")
            return output_path

        filename = Path(file_path).stem
        md_lines = [f"# {filename}\n"]
        headers = rows[0]
        data_rows = rows[1:]
        md_lines.append(f"\n## Data Overview\n")
        md_lines.append(f"- Columns: {len(headers)}")
        md_lines.append(f"- Rows: {len(data_rows)}\n")

        display_cols = min(len(headers), 10)
        display_headers = headers[:display_cols]
        md_lines.append("| " + " | ".join(h.strip() for h in display_headers) + " |")
        md_lines.append("| " + " | ".join("---" for _ in display_headers) + " |")

        for row in data_rows[:100]:
            cells = []
            for i in range(display_cols):
                val = row[i].strip() if i < len(row) else ""
                val = val.replace("|", "\\|").replace("\n", " ")
                if len(val) > 60: val = val[:57] + "..."
                cells.append(val)
            md_lines.append("| " + " | ".join(cells) + " |")

        if len(headers) > 3:
            md_lines.append(f"\n## Record Details\n")
            for i, row in enumerate(data_rows[:50]):
                title = row[0].strip() if row else f"Row {i + 1}"
                if not title or len(title) > 80: title = f"Row {i + 1}"
                md_lines.append(f"\n### {title}\n")
                for j, header in enumerate(headers):
                    val = row[j].strip() if j < len(row) else ""
                    if val: md_lines.append(f"- **{header.strip()}**: {val}")

        result = "\n".join(md_lines)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(result)
        return output_path

    @staticmethod
    def docx_to_markdown(file_path: str, output_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        filename = Path(file_path).stem
        md_lines = [f"# {filename}\n"]

        def _extract_rich_text(para) -> str:
            parts = []
            for run in para.runs:
                text = run.text
                if not text: continue
                if run.bold and run.italic: parts.append(f"***{text}***")
                elif run.bold: parts.append(f"**{text}**")
                elif run.italic: parts.append(f"*{text}*")
                else: parts.append(text)
            return "".join(parts) if parts else para.text

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para = next((p for p in doc.paragraphs if p._element is element), None)
                if not para: continue
                style_name = para.style.name if para.style else ""
                text = para.text.strip()
                if not text:
                    md_lines.append("")
                    continue
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").replace(" ", ""))
                        level = min(level + 1, 6)
                    except ValueError: level = 2
                    md_lines.append(f'\n{"#" * level} {text}\n')
                elif style_name == "Title":
                    md_lines.append(f"\n## {text}\n")
                elif style_name.startswith("List"):
                    md_lines.append(f"- {text}")
                else:
                    md_lines.append(_extract_rich_text(para))
            elif tag == "tbl":
                table = next((t for t in doc.tables if t._element is element), None)
                if not table or not table.rows: continue
                rows = table.rows
                header_cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in rows[0].cells]
                md_lines.append("\n| " + " | ".join(header_cells) + " |")
                md_lines.append("| " + " | ".join("---" for _ in header_cells) + " |")
                for row in rows[1:]:
                    cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
                    md_lines.append("| " + " | ".join(cells) + " |")
                md_lines.append("")

        result = "\n".join(md_lines)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(result)
        return output_path

    @classmethod
    def convert(cls, file_path: str, output_path: str, file_type: str) -> str:
        if file_type == "text": return cls.txt_to_markdown(file_path, output_path)
        elif file_type == "json": return cls.json_to_markdown(file_path, output_path)
        elif file_type == "csv": return cls.csv_to_markdown(file_path, output_path)
        elif file_type == "word": return cls.docx_to_markdown(file_path, output_path)
        else: raise ValueError(f"Unsupported file type: {file_type}")

# Supported extensions and their types
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".json": "json",
    ".csv": "csv",
    ".docx": "word",
    ".doc": "word",
}

def get_file_type(ext: str) -> Optional[str]:
    """Return the file type for a given extension, or None if unsupported."""
    return SUPPORTED_EXTENSIONS.get(ext.lower())
